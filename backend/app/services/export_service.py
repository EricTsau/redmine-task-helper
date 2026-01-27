import io
import re
import json
import base64
import httpx
import markdown
from pathlib import Path
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from urllib.parse import urlparse
from bs4 import BeautifulSoup, NavigableString, Tag
import matplotlib
import matplotlib.pyplot as plt
import numpy as np

# Use non-interactive backend for matplotlib
matplotlib.use('Agg')

class ExportService:
    def __init__(self):
        # Configuration for fonts
        self.font_path = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
        self.bold_font_path = "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"
    
    async def _fetch_image(self, url: str, api_key: str = None, redmine_url: str = None) -> bytes:
        """Helper to fetch image bytes, handling Redmine auth if needed."""
        # Handle local temp images
        if url.startswith("/temp_images/"):
            try:
                # Resolve backend/temp_files directory
                base_dir = Path(__file__).resolve().parent.parent.parent
                filename = url.replace("/temp_images/", "")
                file_path = base_dir / "temp_files" / filename
                if file_path.exists():
                    return file_path.read_bytes()
            except Exception as e:
                print(f"Failed to read local image {url}: {e}")
            return None

        # Handle Data URIs
        if url.startswith("data:image"):
            try:
                # data:image/png;base64,.....
                header, encoded = url.split(",", 1)
                return base64.b64decode(encoded)
            except Exception as e:
                print(f"Failed to decode data URI: {e}")
                return None

        headers = {}
        
        # Check if URL matches Redmine URL
        if redmine_url and api_key and url.startswith("http"):
             # loose check: if url contains redmine domain
             try:
                 r_domain = urlparse(redmine_url).netloc
                 img_domain = urlparse(url).netloc
                 if r_domain and img_domain and r_domain == img_domain:
                     headers["X-Redmine-API-Key"] = api_key
             except:
                 pass
        
        if url.startswith("http"):
            try:
                async with httpx.AsyncClient() as client:
                    resp = await client.get(url, headers=headers, timeout=30.0)
                    if resp.status_code == 200:
                        return resp.content
            except Exception as e:
                print(f"Failed to fetch image {url}: {e}")
        
        return None

    def _generate_gitlab_heatmap_img(self, heatmap_data: list) -> str:
        """Generates a heatmap image and returns it as a base64 data URI."""
        if not heatmap_data:
            return None

        # Prepare data (7 days x 24 hours)
        # heatmap_data is usually [ {day: 0-6, hour: 0-23, count: N}, ... ]
        # We need a 7x24 grid.
        grid = np.zeros((7, 24))
        
        # If input data is structured differently, adjust here.
        # Assuming frontend format: array of { day: int, hour: int, count: int }
        # Or simpler list of list?
        # Based on frontend ActivityHeatmap.tsx, it likely receives processed data.
        # But Report.gitlab_metrics stores raw JSON. 
        # Let's assume the passed `heatmap_data` is compatible or specific format.
        # The frontend `ActivityHeatmap` takes `data[]`.
        
        # Let's try to handle standard list of objects
        try:
            for point in heatmap_data:
                d = int(point.get('day', 0)) # 0=Sun, 6=Sat
                h = int(point.get('hour', 0))
                c = int(point.get('count', 0))
                if 0 <= d < 7 and 0 <= h < 24:
                    grid[d, h] += c
        except Exception as e:
            print(f"Error parsing heatmap data: {e}")
            return None

        days = ["Sun", "Mon", "Tue", "Wed", "Thu", "Fri", "Sat"]
        hours = [str(i) for i in range(24)]

        fig, ax = plt.subplots(figsize=(10, 3))
        # Plot heatmap
        im = ax.imshow(grid, cmap="Greens")
        
        # Axis labels
        ax.set_xticks(np.arange(len(hours)))
        ax.set_yticks(np.arange(len(days)))
        ax.set_xticklabels(hours, fontsize=8)
        ax.set_yticklabels(days, fontsize=8)
        
        plt.setp(ax.get_xticklabels(), rotation=0, ha="center", rotation_mode="anchor")
        
        ax.set_title("Commit Activity Heatmap")
        fig.tight_layout()
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=100)
        plt.close(fig)
        buf.seek(0)
        b64 = base64.b64encode(buf.read()).decode('utf-8')
        return f"data:image/png;base64,{b64}"

    def _generate_gitlab_section(self, gitlab_metrics_json: str) -> str:
        """Parses gitlab_metrics JSON and returns HTML section."""
        if not gitlab_metrics_json:
            return ""
        
        try:
            metrics = json.loads(gitlab_metrics_json)
        except:
            return ""
        
        if not metrics.get('instances'):
            return ""

        html = "<br><hr><h1>GitLab Pulse Dashboard</h1>"
        
        for inst in metrics['instances']:
            name = inst.get('name', 'GitLab')
            impact = inst.get('impact', {})
            cycle = inst.get('cycle', {})
            
            html += f"<h2>Instance: {name}</h2>"
            
            # Summary Table
            html += """<table border="1" width="100%" cellpadding="5">
            <tr>
                <th>Total Commits</th>
                <th>Avg MR Time (h)</th>
                <th>Reviews</th>
                <th>Merged</th>
                <th>Opened</th>
            </tr>
            <tr>
            """
            avg_time = round(cycle.get('average_cycle_time_seconds', 0) / 3600, 1)
            html += f"""
                <td align="center">{impact.get('total_commits', 0)}</td>
                <td align="center">{avg_time}</td>
                <td align="center">{cycle.get('total_review_notes', 0)}</td>
                <td align="center">{cycle.get('merged_count', 0)}</td>
                <td align="center">{cycle.get('opened_count', 0)}</td>
            </tr>
            </table><br>
            """
            
            # Tech Stack
            if impact.get('tech_stack'):
                html += "<h3>Tech Stack</h3>"
                stack_str = ", ".join([f"{t['language']} ({t['percentage']}%)" for t in impact['tech_stack']])
                html += f"<p>{stack_str}</p>"

            # Additions/Deletions
            html += f"<p><b>Additions:</b> +{impact.get('additions', 0)} &nbsp;&nbsp; <b>Deletions:</b> -{impact.get('deletions', 0)}</p>"

            # Heatmap
            if inst.get('heatmap'):
                heatmap_img = self._generate_gitlab_heatmap_img(inst['heatmap'])
                if heatmap_img:
                    html += f'<img src="{heatmap_img}" width="600" alt="Activity Heatmap"><br>'
            
            html += "<br>"
            
        return html

    async def export_to_pdf(self, markdown_text: str, title: str = "Report", api_key: str = None, redmine_url: str = None, gitlab_metrics: str = None) -> io.BytesIO:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Font Loading
        font_loaded = False
        try:
            if Path(self.font_path).exists():
                pdf.add_font("NotoSansCJK", style="", fname=self.font_path) 
                pdf.add_font("NotoSansCJK", style="B", fname=self.bold_font_path)
                font_loaded = True
        except:
             # Try without bold
            try:
                if Path(self.font_path).exists():
                    pdf.add_font("NotoSansCJK", style="", fname=self.font_path)
                    font_loaded = True
            except:
                pass

        if font_loaded:
            pdf.set_font("NotoSansCJK", size=12)
        else:
            pdf.set_font("Helvetica", size=12)
        
        # Convert Markdown to HTML
        html_content = markdown.markdown(markdown_text, extensions=['tables', 'fenced_code'])
        
        # Append GitLab Section
        gitlab_html = self._generate_gitlab_section(gitlab_metrics)
        if gitlab_html:
            html_content += gitlab_html
        
        # Improve table visibility
        html_content = html_content.replace("<table>", '<table border="1" width="100%">')
        
        # Process Images in HTML (Resolve Src)
        # We need to find all srcs first to fetch them async
        soup = BeautifulSoup(html_content, 'html.parser')
        imgs = soup.find_all('img')
        
        # Prefetch logic
        # To avoid blocking loop, we can just fetch sequentially or parallel.
        # Since we use parsing now, simple iteration is fine.
        
        for img in imgs:
            src = img.get('src')
            if not src:
                continue
                
            img_data = await self._fetch_image(src, api_key, redmine_url)
            
            if img_data:
                # Provide base64
                mime = "image/jpeg"
                if src.lower().endswith(".png"): mime = "image/png"
                elif src.lower().endswith(".gif"): mime = "image/gif"
                
                b64_data = base64.b64encode(img_data).decode('utf-8')
                data_uri = f"data:{mime};base64,{b64_data}"
                img['src'] = data_uri
                # Force width for PDF
                img['width'] = "500"
            else:
                # Placeholder
                empty_pixel = "data:image/gif;base64,R0lGODlhAQABAIAAAAAAAP///yH5BAEAAAAALAAAAAABAAEAAAIBRAA7"
                img['src'] = empty_pixel
                img['width'] = "1"
                img['height'] = "1"
        
        # Serialize back to str
        final_html = str(soup)

        # Add Title via HTML too
        full_html = f"<h1>{title}</h1><br>{final_html}"
        
        # Render HTML
        try:
            pdf.write_html(full_html)
        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"PDF write_html error: {e}")
            pdf.cell(0, 10, f"Error rendering PDF content: {str(e)}")
            pdf.ln()
            # fallback raw text not ideal if HTML failed mid-way, but okay
            # pdf.multi_cell(0, 10, markdown_text)

        output = io.BytesIO()
        pdf.output(output)
        output.seek(0)
        return output

    async def export_to_docx(self, markdown_text: str, title: str = "Report", api_key: str = None, redmine_url: str = None, gitlab_metrics: str = None) -> io.BytesIO:
        doc = Document()
        doc.add_heading(title, 0)
        
        # Convert Markdown to HTML
        html_content = markdown.markdown(markdown_text, extensions=['tables', 'fenced_code'])
        
        # Append GitLab Section
        gitlab_html = self._generate_gitlab_section(gitlab_metrics)
        if gitlab_html:
            html_content += gitlab_html

        # Parse HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Recursive traversal helper or flat iteration?
        # Markdown produces relatively flat structure usually, but lists and tables are nested.
        # It's better to iterate over child elements of body (or soup if no body)
        
        elements = soup.find_all(recursive=False)
        if not elements:
            # Maybe it didn't have a body tag, so find_all on soup
            elements = soup.contents
            
        # Recursive function to handle elements
        async def process_element(element, parent_container=None):
            if isinstance(element, NavigableString):
                text = str(element).strip()
                if text and parent_container:
                   parent_container.add_run(text + " ")
                elif text:
                   doc.add_paragraph(text)
                return

            if isinstance(element, Tag):
                tag = element.name
                
                if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    level = int(tag[1])
                    # Python-docx headings are 1-9. 1 is Title-ish logic maybe? default add_heading is level 1
                    # doc.add_heading(text, level)
                    doc.add_heading(element.get_text(), level=level)
                
                elif tag == 'p':
                    p = doc.add_paragraph()
                    # Process children (bold, italic, links within p)
                    await process_children(element, p)
                
                elif tag in ['ul', 'ol']:
                    # Lists
                    style = 'List Bullet' if tag == 'ul' else 'List Number'
                    for li in element.find_all('li', recursive=False):
                        p = doc.add_paragraph(style=style)
                        await process_children(li, p)
                
                elif tag == 'table':
                    rows = element.find_all('tr')
                    if not rows: return
                    # Determine columns
                    cols = len(rows[0].find_all(['td', 'th']))
                    table = doc.add_table(rows=len(rows), cols=cols)
                    table.style = 'Table Grid'
                    
                    for i, row in enumerate(rows):
                        cells = row.find_all(['td', 'th'])
                        for j, cell in enumerate(cells):
                            if j < cols:
                                cell_p = table.cell(i, j).paragraphs[0]
                                await process_children(cell, cell_p)

                elif tag == 'img':
                    src = element.get('src')
                    if src:
                        img_data = await self._fetch_image(src, api_key, redmine_url)
                        if img_data:
                            try:
                                doc.add_picture(io.BytesIO(img_data), width=Inches(6.0))
                            except Exception as e:
                                doc.add_paragraph(f"[Image Error: {e}]")
                        else:
                             doc.add_paragraph(f"[Image Missing: {src}]")
                
                elif tag == 'hr':
                    doc.add_page_break()
                    
                else:
                    # Generic container or unknown, process children
                    # If it's a block level, technically should trigger new paragraph?
                    # But for now, just recurse
                     await process_children(element, parent_container)

        async def process_children(element, paragraph_obj):
            for child in element.contents:
                if isinstance(child, NavigableString):
                    text = str(child)
                    if text:
                        run = paragraph_obj.add_run(text)
                elif isinstance(child, Tag):
                    if child.name == 'strong' or child.name == 'b':
                        run = paragraph_obj.add_run(child.get_text())
                        run.bold = True
                    elif child.name == 'em' or child.name == 'i':
                        run = paragraph_obj.add_run(child.get_text())
                        run.italic = True
                    elif child.name == 'code':
                        run = paragraph_obj.add_run(child.get_text())
                        run.font.name = 'Courier New'
                    elif child.name == 'br':
                        paragraph_obj.add_run("\n")
                    elif child.name == 'a':
                        # Ideally add link, but complicated in python-docx
                        # Just add text styled blue
                        run = paragraph_obj.add_run(child.get_text())
                        run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
                        run.underline = True
                    else:
                        # Recurse? e.g. span
                        # Simply get new text runs
                         run = paragraph_obj.add_run(child.get_text())

        # Start processing
        # Filter top level newlines
        for child in soup.contents:
             await process_element(child)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output
